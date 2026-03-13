import os
import io
import re
import shutil
import tempfile
import logging
import traceback
import redis
from config import settings
from s3_client import s3_client
import json

logger = logging.getLogger(__name__)

redis_client = redis.Redis(
    host=settings.REDIS_HOST,
    port=settings.REDIS_PORT,
    db=settings.REDIS_DB,
    decode_responses=True
)

# ── Constants ───────────────────────────────────────────────────────────────

PT_TO_EMU = 12700  # 1 pt = 12700 EMU (English Metric Units used by python-docx)

# Font‑name mapping: PDF internal name → common Word font name
_FONT_MAP: dict[str, str] = {
    "ArialMT": "Arial",
    "Arial-BoldMT": "Arial",
    "Arial-ItalicMT": "Arial",
    "Arial-BoldItalicMT": "Arial",
    "TimesNewRomanPSMT": "Times New Roman",
    "TimesNewRomanPS-BoldMT": "Times New Roman",
    "TimesNewRomanPS-ItalicMT": "Times New Roman",
    "TimesNewRomanPS-BoldItalicMT": "Times New Roman",
    "Helvetica": "Arial",
    "Helvetica-Bold": "Arial",
    "Helvetica-Oblique": "Arial",
    "Helvetica-BoldOblique": "Arial",
    "CourierNewPSMT": "Courier New",
    "CourierNewPS-BoldMT": "Courier New",
    "Courier": "Courier New",
    "Courier-Bold": "Courier New",
    "Calibri": "Calibri",
    "Calibri-Bold": "Calibri",
    "Calibri-Italic": "Calibri",
    "Calibri-BoldItalic": "Calibri",
    "Cambria": "Cambria",
    "CambriaMath": "Cambria Math",
    "Georgia": "Georgia",
    "Georgia-Bold": "Georgia",
    "Verdana": "Verdana",
    "Verdana-Bold": "Verdana",
    "Tahoma": "Tahoma",
    "Tahoma-Bold": "Tahoma",
    "TrebuchetMS": "Trebuchet MS",
    "Symbol": "Symbol",
    "ZapfDingbats": "Wingdings",
}

# Suffixes to strip when doing fuzzy font matching
_FONT_SUFFIXES = re.compile(
    r"(-(Bold|Italic|BoldItalic|Oblique|BoldOblique|Regular|Light|Medium|Semibold|"
    r"SemiBold|ExtraBold|Black|Thin|Condensed|Narrow)|MT|PS|PSMT)$",
    re.IGNORECASE,
)

# Prefix pattern for embedded font subsets (e.g., "BCDFEE+Arial")
_FONT_PREFIX = re.compile(r"^[A-Z]{6}\+")


def _resolve_font(pdf_font_name: str) -> str:
    """Map a PDF font name to the best Word font name."""
    if not pdf_font_name:
        return "Calibri"

    # Strip subset prefix (e.g., "ABCDEF+TimesNewRomanPSMT" → "TimesNewRomanPSMT")
    clean = _FONT_PREFIX.sub("", pdf_font_name)

    # Direct lookup
    if clean in _FONT_MAP:
        return _FONT_MAP[clean]

    # Try stripping suffixes for a base-name match
    base = _FONT_SUFFIXES.sub("", clean)
    if base in _FONT_MAP:
        return _FONT_MAP[base]

    # Common base names that map directly
    base_lower = base.lower()
    common = {
        "arial": "Arial",
        "helvetica": "Arial",
        "times": "Times New Roman",
        "timesnewroman": "Times New Roman",
        "courier": "Courier New",
        "couriernew": "Courier New",
        "calibri": "Calibri",
        "cambria": "Cambria",
        "georgia": "Georgia",
        "verdana": "Verdana",
        "tahoma": "Tahoma",
        "trebuchet": "Trebuchet MS",
        "trebuchetms": "Trebuchet MS",
        "garamond": "Garamond",
        "palatino": "Palatino Linotype",
        "bookman": "Bookman Old Style",
        "comicsans": "Comic Sans MS",
        "comicsansms": "Comic Sans MS",
        "impact": "Impact",
        "lucidaconsole": "Lucida Console",
        "lucidasans": "Lucida Sans",
        "symbol": "Symbol",
    }
    # Remove non-alphanumeric for fuzzy matching
    normalized = re.sub(r"[^a-z0-9]", "", base_lower)
    if normalized in common:
        return common[normalized]

    # Fall back to the cleaned name itself (Word may still have it)
    return clean if clean else "Calibri"


def _int_to_rgb(color_int: int):
    """Convert PyMuPDF integer color (0xRRGGBB) to python-docx RGBColor."""
    from docx.shared import RGBColor
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return RGBColor(r, g, b)


# ── Scanned PDF detection ──────────────────────────────────────────────────

def _is_scanned_pdf(pdf_path: str) -> bool:
    """Returns True if the PDF contains no extractable text (image-only / scanned)."""
    import fitz
    doc = fitz.open(pdf_path)
    total_chars = sum(len(page.get_text("text").strip()) for page in doc)
    doc.close()
    return total_chars < 100


# ── Paragraph grouping ─────────────────────────────────────────────────────

def _group_lines_into_paragraphs(lines: list[dict], font_size_hint: float = 12.0) -> list[list[dict]]:
    """
    Group consecutive lines into logical paragraphs.
    Lines that are close together vertically and have similar left margins
    are assumed to belong to the same paragraph.
    """
    if not lines:
        return []

    paragraphs: list[list[dict]] = []
    current_para: list[dict] = [lines[0]]

    for i in range(1, len(lines)):
        prev_line = lines[i - 1]
        curr_line = lines[i]

        prev_bottom = prev_line["bbox"][3]
        curr_top = curr_line["bbox"][1]
        vertical_gap = curr_top - prev_bottom

        prev_left = prev_line["bbox"][0]
        curr_left = curr_line["bbox"][0]
        left_diff = abs(curr_left - prev_left)

        # Determine the dominant font size from the previous line
        prev_sizes = []
        for span in prev_line.get("spans", []):
            if span["text"].strip():
                prev_sizes.append(span["size"])
        avg_size = sum(prev_sizes) / len(prev_sizes) if prev_sizes else font_size_hint

        # Same paragraph if vertical gap is small and left margins are similar
        max_gap = avg_size * 0.6
        max_left_diff = 25  # allow some indent variation within a paragraph

        if vertical_gap < max_gap and left_diff < max_left_diff:
            current_para.append(curr_line)
        else:
            paragraphs.append(current_para)
            current_para = [curr_line]

    paragraphs.append(current_para)
    return paragraphs


# ── Text-based PDF conversion ──────────────────────────────────────────────

def _convert_text_pdf(pdf_path: str, docx_path: str):
    """
    Converts text-based PDFs using PyMuPDF directly.
    Preserves: bold, italic, superscript, font name, font size, font color,
    text alignment, indentation, images (proportional), tables, and page dimensions.
    """
    import fitz
    from docx import Document
    from docx.shared import Pt, Inches, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    pdf = fitz.open(pdf_path)

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        page_width_pt = page.rect.width
        page_height_pt = page.rect.height

        # ── Set page dimensions to match PDF ────────────────────────────
        if page_num == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        section.page_width = Emu(int(page_width_pt * PT_TO_EMU))
        section.page_height = Emu(int(page_height_pt * PT_TO_EMU))

        # ── Detect content bounds for margin estimation ─────────────────
        blocks = page.get_text("dict", sort=True)["blocks"]

        content_left = page_width_pt
        content_right = 0
        content_top = page_height_pt
        content_bottom = 0

        for block in blocks:
            bbox = block["bbox"]
            content_left = min(content_left, bbox[0])
            content_right = max(content_right, bbox[2])
            content_top = min(content_top, bbox[1])
            content_bottom = max(content_bottom, bbox[3])

        # Calculate margins (with sensible minimums)
        margin_left = max(content_left, 36)  # at least 0.5 inch
        margin_right = max(page_width_pt - content_right, 36)
        margin_top = max(content_top, 36)
        margin_bottom = max(page_height_pt - content_bottom, 36)

        section.left_margin = Emu(int(margin_left * PT_TO_EMU))
        section.right_margin = Emu(int(margin_right * PT_TO_EMU))
        section.top_margin = Emu(int(margin_top * PT_TO_EMU))
        section.bottom_margin = Emu(int(margin_bottom * PT_TO_EMU))

        content_width = page_width_pt - margin_left - margin_right

        # ── Remove default empty paragraph on first page ────────────────
        if page_num == 0:
            for p in list(doc.paragraphs):
                p._p.getparent().remove(p._p)

        # ── Detect table regions to avoid duplicating text ──────────────
        table_rects = []
        try:
            for tbl in page.find_tables():
                table_rects.append(fitz.Rect(tbl.bbox))
        except Exception:
            pass

        def _in_table(bbox) -> bool:
            r = fitz.Rect(bbox)
            return any(tr.intersects(r) for tr in table_rects)

        # ── Process text and image blocks ───────────────────────────────
        prev_block_bottom = None

        for block in blocks:
            if _in_table(block["bbox"]):
                continue

            if block["type"] == 0:  # text block
                block_lines = block.get("lines", [])
                if not block_lines:
                    continue

                # Filter lines with actual text content
                text_lines = []
                for line in block_lines:
                    spans = [s for s in line["spans"] if s["text"].strip()]
                    if spans:
                        text_lines.append(line)

                if not text_lines:
                    continue

                # Add inter-block spacing
                if prev_block_bottom is not None:
                    gap = block["bbox"][1] - prev_block_bottom
                    avg_span_size = 12
                    for line in text_lines[:1]:
                        for span in line.get("spans", []):
                            if span["text"].strip():
                                avg_span_size = span["size"]
                                break
                    if gap > avg_span_size * 1.5:
                        spacer = doc.add_paragraph()
                        spacer_fmt = spacer.paragraph_format
                        spacer_fmt.space_before = Pt(0)
                        spacer_fmt.space_after = Pt(0)
                        spacer.add_run().font.size = Pt(max(2, int(gap * 0.3)))

                # Group lines into paragraphs
                grouped = _group_lines_into_paragraphs(text_lines)

                for para_lines in grouped:
                    para = doc.add_paragraph()
                    para_format = para.paragraph_format
                    para_format.space_before = Pt(0)
                    para_format.space_after = Pt(0)

                    # ── Alignment detection ─────────────────────────────
                    first_line = para_lines[0]
                    first_line_rect = fitz.Rect(first_line["bbox"])
                    left_gap = first_line_rect.x0 - margin_left
                    right_gap = (margin_left + content_width) - first_line_rect.x1

                    ALIGN_TOLERANCE = 15

                    if content_width > 0:
                        # Right-aligned: text far from left, close to right
                        if right_gap < ALIGN_TOLERANCE and left_gap > content_width * 0.3:
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        # Centered: equal gaps
                        elif abs(left_gap - right_gap) < ALIGN_TOLERANCE and left_gap > ALIGN_TOLERANCE:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Left-aligned with indent
                        elif left_gap > ALIGN_TOLERANCE:
                            para_format.left_indent = Emu(int(left_gap * PT_TO_EMU))

                    # ── Line spacing ────────────────────────────────────
                    if len(para_lines) > 1:
                        first_bottom = para_lines[0]["bbox"][3]
                        second_top = para_lines[1]["bbox"][1]
                        line_gap = second_top - first_bottom
                        first_spans = [s for s in para_lines[0]["spans"] if s["text"].strip()]
                        if first_spans:
                            dominant_size = max(s["size"] for s in first_spans)
                            total_line_height = dominant_size + line_gap
                            if dominant_size > 0:
                                line_spacing_ratio = total_line_height / dominant_size
                                if line_spacing_ratio > 1.1:
                                    from docx.shared import Pt as PtSpacing
                                    para_format.line_spacing = Pt(round(total_line_height))

                    # ── Write runs for each span in each line ───────────
                    for line_idx, line in enumerate(para_lines):
                        spans = [s for s in line["spans"] if s["text"].strip()]
                        if not spans:
                            continue

                        # Add space between lines within the same paragraph
                        if line_idx > 0:
                            run = para.add_run(" ")

                        for span in spans:
                            text = span["text"]
                            run = para.add_run(text)

                            # Font name
                            font_name = _resolve_font(span.get("font", ""))
                            run.font.name = font_name
                            # Set East-Asian font for CJK compatibility
                            r_elem = run._element
                            rPr = r_elem.get_or_add_rPr()
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is None:
                                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                                rPr.insert(0, rFonts)
                            rFonts.set(qn("w:eastAsia"), font_name)

                            # Font size
                            run.font.size = Pt(round(span["size"] * 10) / 10)

                            # Bold / Italic / Superscript
                            flags = span.get("flags", 0)
                            run.bold = bool(flags & 16)       # bit 4
                            run.italic = bool(flags & 2)      # bit 1
                            run.font.superscript = bool(flags & 1)  # bit 0

                            # Font color
                            color_val = span.get("color", 0)
                            if color_val and color_val != 0:
                                run.font.color.rgb = _int_to_rgb(color_val)

                prev_block_bottom = block["bbox"][3]

            elif block["type"] == 1:  # image block
                try:
                    img_bytes = block.get("image")
                    if img_bytes:
                        # Proportional image sizing based on original bbox
                        bbox = block["bbox"]
                        img_width_pt = bbox[2] - bbox[0]
                        img_height_pt = bbox[3] - bbox[1]

                        # Limit to content width
                        if img_width_pt > content_width:
                            scale = content_width / img_width_pt
                            img_width_pt *= scale

                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(
                            io.BytesIO(img_bytes),
                            width=Emu(int(img_width_pt * PT_TO_EMU))
                        )
                        prev_block_bottom = block["bbox"][3]
                except Exception as e:
                    logger.warning(f"Failed to embed image on page {page_num + 1}: {e}")

        # ── Tables ──────────────────────────────────────────────────────
        try:
            for tbl in page.find_tables():
                rows = tbl.extract()
                if not rows:
                    continue

                num_rows = len(rows)
                num_cols = max(len(r) for r in rows)
                if num_cols == 0:
                    continue

                word_tbl = doc.add_table(rows=num_rows, cols=num_cols)
                word_tbl.style = "Table Grid"

                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < num_cols:
                            cell = word_tbl.rows[r_idx].cells[c_idx]
                            cell.text = cell_text or ""

                            # Bold header row (first row)
                            if r_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                doc.add_paragraph()  # spacing after table
        except Exception as e:
            logger.warning(f"Table extraction failed on page {page_num + 1}: {e}")

    pdf.close()
    doc.save(docx_path)
    logger.info(f"Text-based PDF conversion completed: {docx_path}")


# ── Scanned PDF (OCR) conversion ───────────────────────────────────────────

def _convert_scanned_pdf(pdf_path: str, docx_path: str):
    """
    OCR-based converter for scanned / image-only PDFs.
    Renders each page at 300 DPI for high quality, embeds the full-page image
    into the DOCX (preserving all visual content), then appends OCR-extracted
    editable text.
    """
    import fitz
    from pdf2image import convert_from_path
    from docx import Document
    from docx.shared import Inches, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Get page dimensions from the PDF
    pdf_doc = fitz.open(pdf_path)
    page_dims = []
    for p in pdf_doc:
        page_dims.append((p.rect.width, p.rect.height))
    pdf_doc.close()

    doc = Document()

    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)

    images = convert_from_path(pdf_path, dpi=300)

    for page_num, image in enumerate(images):
        # Set page dimensions to match original PDF
        if page_num == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        if page_num < len(page_dims):
            pw, ph = page_dims[page_num]
        else:
            pw, ph = 612, 792  # default letter size in pts

        section.page_width = Emu(int(pw * PT_TO_EMU))
        section.page_height = Emu(int(ph * PT_TO_EMU))

        MARGIN_PT = 36  # 0.5 inch
        section.left_margin = Emu(int(MARGIN_PT * PT_TO_EMU))
        section.right_margin = Emu(int(MARGIN_PT * PT_TO_EMU))
        section.top_margin = Emu(int(MARGIN_PT * PT_TO_EMU))
        section.bottom_margin = Emu(int(MARGIN_PT * PT_TO_EMU))

        content_width_pt = pw - MARGIN_PT * 2

        # Embed full page image
        img_bytes = io.BytesIO()
        image.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(img_bytes, width=Emu(int(content_width_pt * PT_TO_EMU)))

        # OCR text
        try:
            import pytesseract
            ocr_data = pytesseract.image_to_data(
                image, output_type=pytesseract.Output.DICT, lang="eng"
            )

            lines: dict = {}
            for i in range(len(ocr_data["text"])):
                word = ocr_data["text"][i].strip()
                if not word or int(ocr_data["conf"][i]) < 0:
                    continue
                key = (ocr_data["block_num"][i], ocr_data["par_num"][i], ocr_data["line_num"][i])
                if key not in lines:
                    lines[key] = {
                        "words": [],
                        "top": ocr_data["top"][i],
                        "block_num": ocr_data["block_num"][i],
                    }
                lines[key]["words"].append(word)

            if lines:
                sep = doc.add_paragraph()
                sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sep_run = sep.add_run("--- Extracted Text ---")
                sep_run.font.size = Pt(8)
                sep_run.font.color.rgb = _int_to_rgb(0x999999)

                sorted_lines = sorted(lines.values(), key=lambda x: (x["block_num"], x["top"]))
                prev_block = None
                for line_data in sorted_lines:
                    text = " ".join(line_data["words"])
                    if not text.strip():
                        continue
                    if prev_block is not None and line_data["block_num"] != prev_block:
                        doc.add_paragraph()
                    para = doc.add_paragraph()
                    para.add_run(text).font.size = Pt(10)
                    prev_block = line_data["block_num"]

        except Exception as e:
            logger.warning(f"OCR failed on page {page_num + 1}: {e}")

    doc.save(docx_path)
    logger.info(f"OCR conversion (image + text) completed: {docx_path}")


# ── Strategy selector ──────────────────────────────────────────────────────

def _convert_pdf_to_docx_best(pdf_path: str, docx_path: str):
    """Selects text-based or OCR conversion strategy automatically."""
    if _is_scanned_pdf(pdf_path):
        logger.info("Detected scanned PDF - using OCR converter")
        _convert_scanned_pdf(pdf_path, docx_path)
    else:
        logger.info("Detected text-based PDF - using direct fitz converter")
        _convert_text_pdf(pdf_path, docx_path)


# ── Task runner ─────────────────────────────────────────────────────────────

def run_conversion_task(task_id: str, pdf_key: str, docx_key: str):
    """
    Core conversion logic. Runs as a FastAPI BackgroundTask.
    Downloads PDF from S3, converts to DOCX, uploads result, updates Redis status.
    """
    temp_dir = None
    pdf_path = None
    docx_path = None

    try:
        redis_client.setex(
            f"task:{task_id}", 86400,
            json.dumps({"status": "PROCESSING", "progress": 10, "task_id": task_id})
        )

        temp_dir = tempfile.mkdtemp()
        pdf_path = os.path.join(temp_dir, os.path.basename(pdf_key))
        docx_path = os.path.join(temp_dir, os.path.basename(docx_key))

        logger.info(f"Downloading PDF from S3: {pdf_key}")
        s3_client.download_file(pdf_key, pdf_path)

        redis_client.setex(
            f"task:{task_id}", 86400,
            json.dumps({"status": "PROCESSING", "progress": 30, "task_id": task_id})
        )

        logger.info(f"Converting: {os.path.basename(pdf_key)}")
        _convert_pdf_to_docx_best(pdf_path, docx_path)

        redis_client.setex(
            f"task:{task_id}", 86400,
            json.dumps({"status": "PROCESSING", "progress": 80, "task_id": task_id})
        )

        logger.info(f"Uploading DOCX to S3: {docx_key}")
        s3_client.upload_file(docx_path, docx_key)

        download_url = s3_client.generate_presigned_download_url(docx_key)

        redis_client.setex(
            f"task:{task_id}", 86400,
            json.dumps({
                "status": "COMPLETED",
                "progress": 100,
                "task_id": task_id,
                "download_url": download_url,
                "docx_key": docx_key
            })
        )
        logger.info(f"Task {task_id} completed successfully")

    except Exception as e:
        logger.error(f"Error in conversion task {task_id}: {e}\n{traceback.format_exc()}")
        redis_client.setex(
            f"task:{task_id}", 86400,
            json.dumps({"status": "FAILED", "error": str(e), "task_id": task_id})
        )

    finally:
        if pdf_path and os.path.exists(pdf_path):
            os.remove(pdf_path)
        if docx_path and os.path.exists(docx_path):
            os.remove(docx_path)
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
